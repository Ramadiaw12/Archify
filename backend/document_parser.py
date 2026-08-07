# backend/document_parser.py
"""
Extraction de texte depuis différents formats de documents.
Supporte : PDF, DOCX, DOC, TXT, MD, RTF

Hiérarchie des parseurs PDF :
  1. pdfplumber  (meilleure qualité)
  2. PyPDF2      (fallback)
  3. Erreur explicite si aucun n'est disponible
"""

import re
from dataclasses import dataclass, field
from pathlib import Path


# Imports conditionnels 

try:
    import pdfplumber
    _PDF_BACKEND = "pdfplumber"
except ImportError:
    try:
        import PyPDF2          # noqa: F401
        _PDF_BACKEND = "pypdf2"
    except ImportError:
        _PDF_BACKEND = None

try:
    from docx import Document as _DocxDocument
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from striprtf.striprtf import rtf_to_text as _rtf_to_text
    _RTF_OK = True
except ImportError:
    _RTF_OK = False


# Modèle de données

@dataclass
class ParsedDocument:
    """Résultat complet de l'extraction d'un document."""
    filename: str          # Nom original du fichier
    file_type: str         # "PDF", "DOCX", "TXT", etc.
    raw_text: str          # Texte brut nettoyé
    page_count: int        # Nombre de pages (1 si non applicable)
    word_count: int        # Nombre de mots dans raw_text
    char_count: int        # Nombre de caractères
    metadata: dict = field(default_factory=dict)  # Auteur, titre, etc.


# Parseur principal 

class DocumentParser:
    """
    Détecte automatiquement le type de fichier et extrait son texte.
    Tous les parseurs retournent (text: str, pages: int, meta: dict).
    """

    def parse(self, file_path: str, original_filename: str = "") -> ParsedDocument:
        """
        Point d'entrée unique.

        Args:
            file_path: chemin absolu vers le fichier temporaire sur disque
            original_filename: nom du fichier tel qu'uploadé par l'utilisateur

        Returns:
            ParsedDocument avec le texte extrait et les métadonnées

        Raises:
            ValueError: extension non supportée
            RuntimeError: parseur manquant (indiquer le pip install)
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        fname = original_filename or path.name

        dispatch = {
            ".pdf":  self._parse_pdf,
            ".docx": self._parse_docx,
            ".doc":  self._parse_docx,
            ".txt":  self._parse_plaintext,
            ".md":   self._parse_plaintext,
            ".rtf":  self._parse_rtf,
        }

        if ext not in dispatch:
            raise ValueError(
                f"Extension « {ext} » non supportée. "
                f"Formats acceptés : {', '.join(dispatch)}"
            )

        text, pages, meta = dispatch[ext](file_path)
        text = self._clean(text)

        return ParsedDocument(
            filename=fname,
            file_type=ext.lstrip(".").upper(),
            raw_text=text,
            page_count=pages,
            word_count=len(text.split()),
            char_count=len(text),
            metadata=meta,
        )

    #  Parseurs 

    def _parse_pdf(self, path: str) -> tuple[str, int, dict]:
        """Extrait le texte page par page depuis un PDF."""

        if _PDF_BACKEND == "pdfplumber":
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    pages_text.append(t)
                text = "\n\n".join(pages_text)
                meta = {
                    "title":    pdf.metadata.get("Title", ""),
                    "author":   pdf.metadata.get("Author", ""),
                    "producer": pdf.metadata.get("Producer", ""),
                }
                return text, len(pdf.pages), meta

        elif _PDF_BACKEND == "pypdf2":
            import PyPDF2
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages_text = [
                    reader.pages[i].extract_text() or ""
                    for i in range(len(reader.pages))
                ]
                info = reader.metadata or {}
                return "\n\n".join(pages_text), len(reader.pages), {
                    "title":  str(info.get("/Title", "")),
                    "author": str(info.get("/Author", "")),
                }

        else:
            raise RuntimeError(
                "Aucun parseur PDF disponible.\n"
                "Installez pdfplumber :  uv add pdfplumber"
            )

    def _parse_docx(self, path: str) -> tuple[str, int, dict]:
        """Extrait le texte d'un document Word (.docx / .doc)."""
        if not _DOCX_OK:
            raise RuntimeError(
                "python-docx non installé.\n"
                "Installez :  uv add python-docx"
            )
        doc = _DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        props = doc.core_properties
        return text, 1, {
            "title":   props.title or "",
            "author":  props.author or "",
            "subject": props.subject or "",
        }

    def _parse_plaintext(self, path: str) -> tuple[str, int, dict]:
        """Lit un fichier TXT ou Markdown."""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return text, 1, {}

    def _parse_rtf(self, path: str) -> tuple[str, int, dict]:
        """Convertit un fichier RTF en texte brut."""
        if not _RTF_OK:
            raise RuntimeError(
                "striprtf non installé.\n"
                "Installez :  uv add striprtf"
            )
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return _rtf_to_text(content), 1, {}

    #  Nettoyage 

    def _clean(self, text: str) -> str:
        """
        Normalise le texte extrait :
        - Supprime les caractères de contrôle
        - Réduit les sauts de ligne multiples
        - Supprime les espaces de fin de ligne
        - Réduit les espaces internes multiples
        """
        # Supprimer les caractères de contrôle sauf \n et \t
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        # Normaliser les sauts de ligne (max 2 consécutifs)
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Supprimer les espaces en fin de ligne
        text = "\n".join(line.rstrip() for line in text.splitlines())
        # Réduire les espaces multiples en un seul
        text = re.sub(r"[ \t]{2,}", " ", text)
        return text.strip()